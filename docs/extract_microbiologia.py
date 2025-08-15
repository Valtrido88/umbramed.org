import docx
import json
import re

# Si no hay encabezados claros, agrupar por saltos grandes en la numeración de preguntas
def detect_exam_split(question_text, last_number):
    match = re.match(r'^(\d+)\.', question_text)
    if match:
        num = int(match.group(1))
        # Si la numeración baja o salta mucho, es probable que sea otro examen
        if last_number is not None and (num < last_number or num - last_number > 50):
            return True, num
        return False, num
    return False, last_number

# Ruta del archivo Word
import os
doc_path = os.path.join(os.path.dirname(__file__), 'Microbiologia examenes OPE.docx')

# Función para detectar año en las primeras preguntas
def detect_year(questions):
    year_re = re.compile(r'(20\d{2}|19\d{2})')
    for q in questions[:5]:
        match = year_re.search(q['question'])
        if match:
            return match.group(0)
    return None

# Cargar el documento
doc = docx.Document(doc_path)



exams = {}
exam_count = 1
current_exam = f"Examen_{exam_count}"
questions = []
current_question = None
options = []
answer = None
last_number = None

# Expresiones regulares para detectar preguntas y respuestas
question_re = re.compile(r'^\d+\..+')
option_re = re.compile(r'^[a-dA-D]\)')
answer_re = re.compile(r'^Respuesta:')



for para in doc.paragraphs:
    text = para.text.strip()
    if question_re.match(text):
        split, num = detect_exam_split(text, last_number)
        if split:
            # Guardar el bloque anterior y empezar uno nuevo
            if current_exam and questions:
                exams[current_exam] = questions
            exam_count += 1
            current_exam = f"Examen_{exam_count}"
            questions = []
        last_number = num
        if current_question:
            questions.append({
                'question': current_question,
                'options': options,
                'answer': answer
            })
        current_question = text
        options = []
        answer = None
    elif option_re.match(text):
        options.append(text)
    elif answer_re.match(text):
        answer = text.replace('Respuesta:', '').strip()


# Añadir la última pregunta y el último bloque de examen
if current_question:
    questions.append({
        'question': current_question,
        'options': options,
        'answer': answer
    })
if current_exam and questions:
    exams[current_exam] = questions

# Guardar como JSON



# Extraer también desde tablas, agrupando por examen si hay saltos grandes en la numeración
for table in doc.tables:
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) >= 1 and question_re.match(cells[0]):
            split, num = detect_exam_split(cells[0], last_number)
            if split:
                if current_exam and questions:
                    exams[current_exam] = questions
                exam_count += 1
                current_exam = f"Examen_{exam_count}"
                questions = []
            last_number = num
            q_text = cells[0]
            opts = cells[1:-1] if len(cells) > 2 else []
            ans = cells[-1] if len(cells) > 1 else None
            if q_text:
                questions.append({
                    'question': q_text,
                    'options': opts,
                    'answer': ans
                })
    if current_exam and questions:
        exams[current_exam] = questions



# Renombrar los bloques usando el año si se detecta
renamed_exams = {}
for key, questions in exams.items():
    year = detect_year(questions)
    if year:
        new_key = f"OPE_{year}"
    else:
        new_key = key
    renamed_exams[new_key] = questions

output_path = os.path.join(os.path.dirname(__file__), '../data/examen_microbiologia.json')
with open(output_path, 'w', encoding='utf-8') as f:
    json.dump(renamed_exams, f, ensure_ascii=False, indent=2)

print(f'Exámenes extraídos: {len(renamed_exams)}. Guardado en {output_path}')
